# -*- coding: utf-8 -*-
"""Build a Word review draft from the final SSE paper Markdown manuscript."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CONTENT_WIDTH_DXA = 9360
INLINE_TOKEN_RE = re.compile(r"(`[^`]+`|\*\*.*?\*\*|\[[^\]]+\]\([^)]+\))")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        element = tc_mar.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def get_or_add_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = get_or_add_child(tbl_pr, "w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = get_or_add_child(tbl_pr, "w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = get_or_add_child(tc_pr, "w:tcW")
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def column_widths(rows: list[list[str]]) -> list[int]:
    cols = max(len(row) for row in rows)
    weights = [1.0] * cols
    for c in range(cols):
        max_len = max((len(row[c]) if c < len(row) else 0) for row in rows)
        weights[c] = max(0.7, min(2.8, max_len / 12))
    total = sum(weights)
    widths = [max(650, int(CONTENT_WIDTH_DXA * w / total)) for w in weights]
    delta = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += delta
    return widths


def add_inline_runs(para, text: str, size_pt: float = 10.5, bold: bool = False) -> None:
    for part in INLINE_TOKEN_RE.split(text):
        if not part:
            continue
        link = LINK_RE.fullmatch(part)
        if part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            set_run_font(run, max(size_pt - 1, 8.0), bold)
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        elif part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            set_run_font(run, size_pt, True)
        elif link:
            run = para.add_run(f"{link.group(1)} ({link.group(2)})")
            set_run_font(run, size_pt, bold)
        else:
            run = para.add_run(part)
            set_run_font(run, size_pt, bold)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 14, 8),
        ("Heading 2", 13, "2E74B5", 10, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_formatted_paragraph(doc: Document, text: str) -> None:
    stripped = text.strip()
    if not stripped:
        return
    if stripped.startswith("> "):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(8)
        run = para.add_run(stripped[2:])
        set_run_font(run, 10)
        run.italic = True
        return
    if stripped.startswith("- "):
        para = doc.add_paragraph(style="List Bullet")
        add_inline_runs(para, stripped[2:], 10.5)
        return
    if re.match(r"^\d+\. ", stripped):
        para = doc.add_paragraph(style="List Number")
        add_inline_runs(para, re.sub(r"^\d+\. ", "", stripped), 10.5)
        return
    para = doc.add_paragraph()
    add_inline_runs(para, stripped, 10.5)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = column_widths(rows)
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            text = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            add_inline_runs(para, text, 8.2 if cols >= 6 else 9.2, bold=(r_idx == 0))
            if r_idx == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F2F4F7")
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()


def add_figure(doc: Document, input_md: Path, alt_text: str, image_ref: str) -> None:
    image_path = Path(image_ref)
    if not image_path.is_absolute():
        candidates = [input_md.parent / image_path, Path.cwd() / image_path]
        image_path = next((candidate for candidate in candidates if candidate.exists()), candidates[0])
    if not image_path.exists():
        para = doc.add_paragraph()
        run = para.add_run(f"[Missing figure: {image_ref}]")
        set_run_font(run, 9)
        run.italic = True
        return

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(image_path), width=Inches(6.05))
    if alt_text:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)
        run = caption.add_run(alt_text)
        set_run_font(run, 9)
        run.italic = True


def build_docx(input_md: Path, output_docx: Path) -> None:
    doc = Document()
    style_document(doc)
    lines = input_md.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    first_heading = True
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.25)
                for code_line in code_lines:
                    run = para.add_run(code_line + "\n")
                    run.font.name = "Consolas"
                    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
                    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
                    run.font.size = Pt(8.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table(table_lines))
            continue
        image_match = IMAGE_RE.fullmatch(stripped)
        if image_match:
            add_figure(doc, input_md, image_match.group(1), image_match.group(2))
            i += 1
            continue
        if stripped.startswith("# "):
            text = stripped[2:].strip()
            if first_heading:
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_after = Pt(12)
                run = para.add_run(text)
                set_run_font(run, 20, True)
                run.font.color.rgb = RGBColor.from_string("0B2545")
                first_heading = False
            else:
                doc.add_heading(text, level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:].strip(), level=3)
        else:
            add_formatted_paragraph(doc, line)
        i += 1

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("SSE Future-Slip Forecasting Review Draft")
        set_run_font(run, 8)
        run.font.color.rgb = RGBColor.from_string("666666")

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input-md", default="docs/final_paper_manuscript_zh.md")
    parser.add_argument("--output-docx", default="docs/final_paper_manuscript_zh.docx")
    args = parser.parse_args()
    build_docx(Path(args.input_md), Path(args.output_docx))
    print(args.output_docx)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
